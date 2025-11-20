import sys
import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    os.system(f'"{sys.executable}" -m pip install python-docx')
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title
title = doc.add_heading('AI Coding Assistants - Token Limit Comparison', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Claude Sonnet 4.5 - Comprehensive Analysis')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(102, 126, 234)

doc.add_paragraph()

# Overview
doc.add_heading('🎯 Overview', 1)
doc.add_paragraph('This comparison focuses specifically on how each AI coding assistant handles Claude Sonnet 4.5 and what happens when you hit token limits.')

doc.add_paragraph()

# Tools Compared
doc.add_heading('Tools Compared:', 2)
tools = [
    ('Cursor', 'AI-powered code editor (fork of VS Code)'),
    ('Windsurf', 'AI coding assistant by Codeium (VS Code or standalone)'),
    ('Cline', 'VS Code extension with bring-your-own-key (BYOK) API access'),
    ('GitHub Copilot', "Microsoft's AI coding assistant (multi-IDE support)"),
    ('Claude Code', 'Official Anthropic VS Code extension for AI-assisted coding')
]

for i, (name, desc) in enumerate(tools, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(f'{name}').bold = True
    p.add_run(f' - {desc}')

doc.add_page_break()

# Quick Comparison Table
doc.add_heading('📊 Quick Comparison Table', 1)

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'

# Header row
headers = ['Tool', 'Max Context', 'Auto-Refresh', 'Context Preservation', 'Cost Model', 'At Limit']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

# Data rows
data = [
    ['Cursor', '1,000,000 tokens\n(Sonnet 4.5)', '✅ Yes', 'Excellent', 'Subscription\n$20/mo', 'Seamless continuation'],
    ['Windsurf', '1,000,000 tokens\n(Sonnet 4.5)', '✅ Yes', 'Excellent', 'Subscription\n$10/mo', 'Seamless continuation'],
    ['Cline', '200,000 tokens\n(Sonnet 4.5)', '❌ No', 'Manual', 'BYOK (API)\n~$18/mo', 'Hard stop, new task'],
    ['GitHub Copilot', '200,000 tokens\n(Sonnet 4)', '❌ No', 'Manual', 'Subscription\n$10/mo', 'Hard stop, new chat'],
    ['Claude Code', '200,000 tokens\n(Sonnet 4.5)', '❌ No', 'Manual', 'Subscription\n$20/mo', 'Hard stop, new conversation']
]

for i, row_data in enumerate(data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_page_break()

# Detailed Analysis - Cursor
doc.add_heading('1️⃣ Cursor', 1)

doc.add_heading('Token Limits', 2)
p = doc.add_paragraph()
p.add_run('Model: ').bold = True
p.add_run('Claude Sonnet 4.5\n')
p.add_run('Context Window: ').bold = True
p.add_run('1,000,000 tokens\n')
p.add_run('Effective Usage: ').bold = True
p.add_run('~950,000 tokens (buffer for responses)')

doc.add_heading('What Happens at Token Limit', 2)
doc.add_paragraph('When approaching 1M tokens:')
features = [
    '🔄 Automatic refresh triggered',
    '✅ TODO list items preserved',
    '✅ Current task context maintained',
    '✅ Key architectural decisions saved',
    '✅ Files being worked on tracked',
    '✅ Recent changes summarized',
    '✅ New context window created seamlessly'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Pros & Cons', 2)
doc.add_paragraph('Pros:', style='List Bullet').runs[0].bold = True
pros = [
    'Massive 1M token context',
    'Automatic refresh - no interruption',
    'Excellent context preservation',
    'Great for large refactoring projects',
    'Built-in IDE with AI integration'
]
for pro in pros:
    doc.add_paragraph(f'✅ {pro}', style='List Bullet 2')

doc.add_paragraph('Cons:', style='List Bullet').runs[0].bold = True
cons = [
    'Higher cost ($20/month)',
    'Locked to Cursor IDE',
    'Cannot use with existing VS Code setup'
]
for con in cons:
    doc.add_paragraph(f'❌ {con}', style='List Bullet 2')

doc.add_page_break()

# Windsurf
doc.add_heading('2️⃣ Windsurf', 1)

doc.add_heading('Token Limits', 2)
p = doc.add_paragraph()
p.add_run('Model: ').bold = True
p.add_run('Claude Sonnet 4.5\n')
p.add_run('Context Window: ').bold = True
p.add_run('1,000,000 tokens\n')
p.add_run('Effective Usage: ').bold = True
p.add_run('~950,000 tokens')

doc.add_heading('Key Features', 2)
features = [
    '🔄 Automatic refresh like Cursor',
    '🚀 Cascade planning - creates multi-step plans',
    '🤖 Agent mode with autonomous execution',
    '💬 Flow mode for continuous assistance',
    '💰 Best price: $10/month for 1M context'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Pros & Cons', 2)
doc.add_paragraph('Pros:', style='List Bullet').runs[0].bold = True
pros = [
    'Best value: $10/month for 1M tokens',
    'Advanced Cascade planning',
    'Autonomous agent capabilities',
    'Works with VS Code or standalone',
    'Automatic context refresh'
]
for pro in pros:
    doc.add_paragraph(f'✅ {pro}', style='List Bullet 2')

doc.add_paragraph('Cons:', style='List Bullet').runs[0].bold = True
cons = [
    'Newer tool (less mature)',
    'Smaller community than Cursor',
    'Some features still in development'
]
for con in cons:
    doc.add_paragraph(f'❌ {con}', style='List Bullet 2')

doc.add_page_break()

# Cline
doc.add_heading('3️⃣ Cline', 1)

doc.add_heading('Token Limits', 2)
p = doc.add_paragraph()
p.add_run('Model: ').bold = True
p.add_run('Claude Sonnet 4.5\n')
p.add_run('Context Window: ').bold = True
p.add_run('200,000 tokens\n')
p.add_run('Cost Model: ').bold = True
p.add_run('BYOK (Bring Your Own Key) - Pay per usage')

doc.add_heading('What Happens at Token Limit', 2)
doc.add_paragraph('❌ Hard stop at 200k tokens')
doc.add_paragraph('⚠️ Must start new task/conversation')
doc.add_paragraph('📝 Manual context transfer required')

doc.add_heading('Pros & Cons', 2)
doc.add_paragraph('Pros:', style='List Bullet').runs[0].bold = True
pros = [
    'Full control over API usage',
    'Transparent token counting',
    'Pay only for what you use',
    'Works in VS Code',
    'Agent mode with task-based workflow',
    'Great for privacy (self-hosted API possible)'
]
for pro in pros:
    doc.add_paragraph(f'✅ {pro}', style='List Bullet 2')

doc.add_paragraph('Cons:', style='List Bullet').runs[0].bold = True
cons = [
    'No automatic refresh',
    'Smaller context (200k)',
    'Requires API key management',
    'Multiple conversations needed for large tasks'
]
for con in cons:
    doc.add_paragraph(f'❌ {con}', style='List Bullet 2')

doc.add_page_break()

# GitHub Copilot
doc.add_heading('4️⃣ GitHub Copilot', 1)

doc.add_heading('Token Limits', 2)
p = doc.add_paragraph()
p.add_run('Model: ').bold = True
p.add_run('Claude Sonnet 4 (not 4.5)\n')
p.add_run('Context Window: ').bold = True
p.add_run('200,000 tokens\n')
p.add_run('Cost: ').bold = True
p.add_run('$10/month (Individual) or $19/user/month (Business)')

doc.add_heading('Key Features', 2)
features = [
    '✅✅✅ Best-in-class autocomplete',
    '✅ Agent mode via Copilot Chat',
    '✅ Multi-IDE support (VS Code, Visual Studio, JetBrains, etc.)',
    '✅ GitHub integration',
    '✅ Enterprise features available'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Pros & Cons', 2)
doc.add_paragraph('Pros:', style='List Bullet').runs[0].bold = True
pros = [
    'Industry-leading autocomplete',
    'Works in multiple IDEs',
    'Strong GitHub integration',
    'Excellent enterprise support',
    'Agent mode capabilities',
    'Good price ($10/month)'
]
for pro in pros:
    doc.add_paragraph(f'✅ {pro}', style='List Bullet 2')

doc.add_paragraph('Cons:', style='List Bullet').runs[0].bold = True
cons = [
    'Only Claude Sonnet 4 (not 4.5)',
    'Smaller context (200k)',
    'No automatic refresh',
    'Chat is secondary to autocomplete'
]
for con in cons:
    doc.add_paragraph(f'❌ {con}', style='List Bullet 2')

doc.add_page_break()

# Claude Code
doc.add_heading('5️⃣ Claude Code (Official Anthropic VS Code Extension)', 1)

doc.add_heading('Token Limits', 2)
p = doc.add_paragraph()
p.add_run('Model: ').bold = True
p.add_run('Claude Sonnet 4.5 ✅\n')
p.add_run('Context Window: ').bold = True
p.add_run('200,000 tokens\n')
p.add_run('Platform: ').bold = True
p.add_run('VS Code Extension (Official Anthropic)\n')
p.add_run('Cost: ').bold = True
p.add_run('$20/month')

doc.add_heading('Key Features', 2)
features = [
    '✅ Official Anthropic extension',
    '✅ Agent mode with multi-file editing',
    '✅ File system access (read/write)',
    '✅ Workspace awareness',
    '✅ VS Code integration',
    '✅ Latest Claude Sonnet 4.5 model'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('What Happens at Token Limit', 2)
doc.add_paragraph('❌ Hard stop at 200k tokens')
doc.add_paragraph('⚠️ Must start new conversation')
doc.add_paragraph('📝 Manual context summarization required')

doc.add_heading('Pros & Cons', 2)
doc.add_paragraph('Pros:', style='List Bullet').runs[0].bold = True
pros = [
    'Official Anthropic product',
    'Latest Claude Sonnet 4.5',
    'Agent mode capabilities',
    'File system access',
    'Multi-file editing',
    'VS Code integration'
]
for pro in pros:
    doc.add_paragraph(f'✅ {pro}', style='List Bullet 2')

doc.add_paragraph('Cons:', style='List Bullet').runs[0].bold = True
cons = [
    'Smaller context (200k)',
    'No automatic refresh',
    'No terminal execution',
    'Multiple conversations needed for large tasks',
    'Higher cost for limited context'
]
for con in cons:
    doc.add_paragraph(f'❌ {con}', style='List Bullet 2')

doc.add_page_break()

# Feature Comparison Matrix
doc.add_heading('📊 Feature Comparison Matrix', 1)

table = doc.add_table(rows=10, cols=6)
table.style = 'Light Grid Accent 1'

# Headers
headers = ['Feature', 'Cursor', 'Windsurf', 'Cline', 'Copilot', 'Claude Code']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

# Data
features_data = [
    ['Claude Model', 'Sonnet 4.5', 'Sonnet 4.5', 'Sonnet 4.5', 'Sonnet 4', 'Sonnet 4.5'],
    ['Context Size', '1M', '1M', '200k', '200k', '200k'],
    ['Auto-Refresh', '✅', '✅', '❌', '❌', '❌'],
    ['File Access', '✅', '✅', '✅', '✅', '✅'],
    ['Terminal Exec', '✅', '✅', '✅', 'Limited', '❌'],
    ['Autocomplete', '✅', '✅', '❌', '✅✅✅', '❌'],
    ['Agent Mode', '✅', '✅✅', '✅', '✅', '✅'],
    ['IDE', 'Cursor', 'Windsurf/VSC', 'VS Code', 'Multiple', 'VS Code'],
    ['Cost/Month', '$20', '$10', '~$18 BYOK', '$10', '$20']
]

for i, row_data in enumerate(features_data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_page_break()

# Recommendations
doc.add_heading('🏆 Recommendations', 1)

doc.add_heading('Best for Large Projects & Refactoring:', 2)
p = doc.add_paragraph()
p.add_run('Cursor or Windsurf').bold = True
p = doc.add_paragraph('Both offer 1M token context with automatic refresh. Choose Windsurf for best value ($10/mo) or Cursor for more mature ecosystem ($20/mo).')

doc.add_heading('Best for Budget-Conscious Teams:', 2)
p = doc.add_paragraph()
p.add_run('Windsurf').bold = True
p = doc.add_paragraph('$10/month for 1M tokens is unbeatable value. Includes Cascade planning and agent mode.')

doc.add_heading('Best for Autocomplete Focus:', 2)
p = doc.add_paragraph()
p.add_run('GitHub Copilot').bold = True
p = doc.add_paragraph('Industry-leading autocomplete at $10/month. Now includes agent mode. Works across multiple IDEs.')

doc.add_heading('Best for Control & Transparency:', 2)
p = doc.add_paragraph()
p.add_run('Cline').bold = True
p = doc.add_paragraph('BYOK model gives full control. Great for privacy-conscious teams or those wanting detailed usage tracking.')

doc.add_heading('Best for Official Anthropic Support:', 2)
p = doc.add_paragraph()
p.add_run('Claude Code').bold = True
p = doc.add_paragraph('Direct from Anthropic with latest Sonnet 4.5. Best if you want official support and VS Code integration.')

doc.add_page_break()

# Enterprise Comparison
doc.add_heading('🏢 Enterprise Features & Pricing', 1)

table = doc.add_table(rows=6, cols=6)
table.style = 'Light Grid Accent 1'

headers = ['Feature', 'Cursor', 'Windsurf', 'Cline', 'Copilot', 'Claude Code']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

enterprise_data = [
    ['Enterprise Price', '$40/user/mo', '$15/user/mo', 'Custom', '$39/user/mo', '$30/user/mo'],
    ['SSO/SAML', '✅', '✅', 'Self-hosted', '✅', '✅'],
    ['Admin Dashboard', '✅', '✅', '❌', '✅✅', '✅'],
    ['SLA Guarantee', '✅', 'Contact', '❌', '✅✅', '✅'],
    ['Priority Support', '✅', '✅', 'Community', '✅✅', '✅']
]

for i, row_data in enumerate(enterprise_data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_paragraph()
doc.add_paragraph('For 10 developers annually:')
costs = [
    'Windsurf Teams: $1,800/year (Best value)',
    'GitHub Copilot Business: $2,280/year',
    'Cline (BYOK): ~$2,160/year (variable)',
    'Claude Code Team: $3,600/year',
    'Cursor Business: $4,800/year'
]
for cost in costs:
    doc.add_paragraph(cost, style='List Bullet')

doc.add_page_break()

# Conclusion
doc.add_heading('📝 Conclusion', 1)

doc.add_paragraph('All five AI coding assistants now include agent mode capabilities, making them powerful tools for software development. The key differentiators are:')

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Context Size: ').bold = True
p.add_run('Cursor and Windsurf lead with 1M tokens vs 200k for others')

p = doc.add_paragraph()
p.add_run('Auto-Refresh: ').bold = True
p.add_run('Only Cursor and Windsurf offer seamless context continuation')

p = doc.add_paragraph()
p.add_run('Price/Value: ').bold = True
p.add_run('Windsurf offers best value at $10/mo for 1M context')

p = doc.add_paragraph()
p.add_run('Autocomplete: ').bold = True
p.add_run('GitHub Copilot remains the autocomplete champion')

p = doc.add_paragraph()
p.add_run('Official Support: ').bold = True
p.add_run('Claude Code is the official Anthropic VS Code extension')

doc.add_paragraph()
doc.add_paragraph('Choose based on your priorities: context size, budget, IDE preference, or specific features like autocomplete or enterprise support.')

# Footer
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph('_______________________________________________')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('Generated from Flight Booking System Documentation')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
p = doc.add_paragraph('https://github.com/sangyvann/Flight-Booking-System-JavaServlets_App')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(102, 126, 234)

# Save
doc.save('AI-Coding-Assistants-Comparison.docx')
print("✅ Word document created: AI-Coding-Assistants-Comparison.docx")

